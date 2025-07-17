#!/usr/bin/env python3
"""
Ardonie Capital One Page Pitch PDF Generator
Generates a professional PDF version of the one-page business pitch
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
import matplotlib.pyplot as plt
import io
import os
import plotly.graph_objects as go

def create_ardonie_capital_pitch_pdf():
    """Generate Ardonie Capital one-page pitch PDF"""

    # Create document
    doc = Document()
    section = doc.sections[0]
    section.page_height, section.page_width = section.page_width, section.page_height
    section.orientation = WD_ORIENT.LANDSCAPE

    # Title
    title = doc.add_heading('Ardonie Capital: One Page Business Pitch', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact info
    contact_p = doc.add_paragraph('ardoniecapital.com | DFW Auto Repair Shop Marketplace | Express Deal Package - 34 Days')
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_highlighted(paragraph, text, color):
        """Add highlighted text to paragraph"""
        run = paragraph.add_run(text)
        run.font.bold = True
        run.font.color.rgb = color

    # Section: The Problem
    p = doc.add_paragraph()
    add_highlighted(p, 'The Problem: ', RGBColor(255, 0, 0))
    p.add_run('Auto repair shop owners in the DFW area face significant challenges when buying or selling their businesses. Traditional transactions take 6-12 months, involve complex paperwork, and lack specialized industry expertise. Buyers struggle to find quality shops, while sellers can\'t efficiently reach qualified buyers. The fragmented process involves disconnected legal, financial, and broker services.')

    # Section: The Solution
    p = doc.add_paragraph()
    add_highlighted(p, 'The Solution: ', RGBColor(0, 112, 192))
    p.add_run('Ardonie Capital is the premier marketplace specifically designed for DFW auto repair shop transactions. Our Express Deal Package reduces transaction time to just 34 days through pre-vetted networks of automotive industry specialists, streamlined documentation, Express Buyer and Express Seller badge systems, dedicated deal rooms with milestone tracking, and integrated funding assistance.')

    # Section: Value Proposition
    p = doc.add_paragraph()
    add_highlighted(p, 'Value Proposition: ', RGBColor(0, 176, 80))
    p.add_run('We make buying or selling an auto repair shop as efficient as our Express Deal Package promises. Sellers get expert automotive industry guidance and access to qualified buyers. Buyers enjoy streamlined processes, transparent due diligence, and built-in access to specialized financing. Legal firms and financial institutions receive a pipeline of qualified, ready-to-transact automotive business clients.')

    # Section: Target Market
    p = doc.add_paragraph()
    add_highlighted(p, 'Target Market: ', RGBColor(112, 48, 160))
    p.add_run('2,500+ independent auto repair shops in DFW with $2.1B total market value. 15% annual ownership turnover rate creates 375+ transactions annually. Average shop value ranges from $200K - $2M. Growing consolidation trend in automotive services creates expanding market opportunity.')

    # Section: Business Model
    p = doc.add_paragraph()
    add_highlighted(p, 'Business Model: ', RGBColor(255, 192, 0))
    p.add_run('Revenue streams include 3-5% transaction fees on successful deals, Express Badge premium memberships, professional referral fees from legal and financial partners, and premium listing services. Projected Year 1 revenue: $500K-$800K from 25-40 transactions.')

    # Section: Competitive Advantage
    p = doc.add_paragraph()
    add_highlighted(p, 'Competitive Advantage: ', RGBColor(0, 176, 240))
    p.add_run('Industry Specialization: First platform dedicated exclusively to auto repair shop transactions. Geographic Focus: Deep DFW market knowledge and local partnerships. Speed to Close: 34-day Express Deal Package vs. 6-12 month industry standard. Vetted Network: Pre-qualified professionals and streamlined processes.')

    # Section: Market Opportunity
    p = doc.add_paragraph()
    add_highlighted(p, 'Market Opportunity: ', RGBColor(146, 208, 80))
    p.add_run('$2.1B annual market value in DFW alone. 375+ annual transactions with average $5K-$25K in legal fees per transaction. Growing demand for acquisition financing and specialized automotive industry expertise. Expansion opportunities to Houston, Austin, San Antonio, and beyond.')

    # Section: Financial Projections
    p = doc.add_paragraph()
    add_highlighted(p, 'Financial Projections: ', RGBColor(255, 0, 102))
    p.add_run('Year 1: $500K-$800K revenue, 25-40 transactions. Year 2: $1.2M-$2M revenue, 60-100 transactions. Year 3: $2M-$3.5M revenue, 100-150 transactions. Target 25% market share in DFW by Year 3.')

    # Section: The Ask
    p = doc.add_paragraph()
    add_highlighted(p, 'The Ask: ', RGBColor(0, 32, 96))
    p.add_run('We\'re seeking strategic partners and investors to help scale our Express Deal Package platform. Join our network of preferred legal firms, financial institutions, and industry professionals. Become part of the future of DFW auto repair shop transactions.')

    # Section: Call to Action
    p = doc.add_paragraph()
    add_highlighted(p, 'Call to Action: ', RGBColor(255, 102, 0))
    p.add_run('Ready to transform DFW auto repair shop transactions? Contact Ardonie Capital to discuss partnership opportunities, investment, or to list your shop today.')

    # Tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_highlighted(p, 'DFW Auto Repair Deals. Done in 34 Days.', RGBColor(0, 112, 192))

    # Add a pie chart for Market Breakdown
    try:
        labels = ['Auto Repair Shops', 'Transmission Shops', 'Specialty Services']
        sizes = [60, 25, 15]
        colors = ['#0070C0', '#92D050', '#FFC000']
        fig, ax = plt.subplots(figsize=(3, 3))
        ax.pie(sizes, labels=labels, colors=colors, autopct='%1.0f%%', startangle=140)
        ax.set_title('DFW Auto Service Market Breakdown')
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(2.5))
    except Exception as e:
        print(f"Could not generate pie chart: {e}")

    # Add a bar chart for Deal Timeline Comparison
    try:
        labels = ['Industry Avg', 'Ardonie Capital\nExpress Deal']
        values = [270, 34]  # 6-12 months avg vs 34 days
        fig, ax = plt.subplots(figsize=(3, 2))
        bars = ax.bar(labels, values, color=['#C00000', '#0070C0'])
        ax.set_ylabel('Days to Close')
        ax.set_title('Deal Timeline: Industry vs. Ardonie Capital')

        # Add value labels on bars
        for i, (bar, v) in enumerate(zip(bars, values)):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 5,
                   f'{v} days', ha='center', va='bottom', fontweight='bold')

        img_stream2 = io.BytesIO()
        plt.savefig(img_stream2, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        img_stream2.seek(0)
        doc.add_picture(img_stream2, width=Inches(2.5))
    except Exception as e:
        print(f"Could not generate bar chart: {e}")

    # Create documents directory if it doesn't exist
    os.makedirs('../documents', exist_ok=True)

    # Save the document
    docx_path = '../documents/Ardonie_Capital_One_Page_Pitch.docx'
    doc.save(docx_path)

    print(f"Generated: {docx_path}")

    # Try to convert to PDF if possible
    try:
        # Try using docx2pdf if available
        import docx2pdf
        pdf_path = '../documents/Ardonie_Capital_One_Page_Pitch.pdf'
        docx2pdf.convert(docx_path, pdf_path)
        print(f"Generated PDF: {pdf_path}")
        return pdf_path
    except ImportError:
        print("docx2pdf not available. DOCX file generated successfully.")
        print("To convert to PDF, install: pip install docx2pdf")
        return docx_path
    except Exception as e:
        print(f"PDF conversion failed: {e}")
        return docx_path

if __name__ == "__main__":
    create_ardonie_capital_pitch_pdf()
